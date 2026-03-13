# delta_cv_generator.py
# Style-driven Word CV generator for Delta
# Uses template from Knowledge and swaps ONLY body profile image (image1.png)
#
# Fixes:
# - Normalizes mapping keys so both "NAME" and "{{NAME}}" work (prevents leaving {{ }} in output)
# - Replaces placeholders in body + headers/footers
# - Applies NAME formatting (bold + 16pt) during replacement (so it still works after token replacement)

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from zipfile import ZipFile
from PIL import Image
import io
import os
import re

PROFILE_PHOTO_MEDIA_PATH = "word/media/image1.png"
PHOTO_TARGET_PX = 900


# -----------------------------
# SAFE TEXT HELPERS
# -----------------------------
def paragraph_has_drawing(p: Paragraph) -> bool:
    xml = p._element.xml
    return "<w:drawing" in xml or "<wp:anchor" in xml or "<wp:inline" in xml


def set_first_text_node_in_paragraph(p: Paragraph, new_text: str):
    p_elm = p._element
    t_el = p_elm.find(".//w:t", namespaces=p_elm.nsmap)
    if t_el is not None:
        t_el.text = new_text
        for extra in p_elm.findall(".//w:t", namespaces=p_elm.nsmap)[1:]:
            extra.text = ""


def remove_all_runs(p: Paragraph):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)


def set_paragraph_text(p: Paragraph, text: str, bold=None, size_pt=None):
    """Replaces all visible text in the paragraph while preserving drawings if present."""
    if paragraph_has_drawing(p):
        set_first_text_node_in_paragraph(p, text)
        return

    remove_all_runs(p)
    run = p.add_run(text)
    run.bold = bold
    if size_pt is not None:
        run.font.size = Pt(size_pt)


# -----------------------------
# BULLETS
# -----------------------------
def insert_paragraph_after(doc: Document, after_p: Paragraph, style):
    new_p = OxmlElement("w:p")
    after_p._element.addnext(new_p)
    p = Paragraph(new_p, doc._body)
    p.style = style
    return p


def delete_paragraph(p: Paragraph):
    p._element.getparent().remove(p._element)


def expand_bullets(doc: Document, placeholder: str, bullets: list):
    target = None
    for p in doc.paragraphs:
        if placeholder in p.text:
            target = p
            break
    if not target:
        return

    style = target.style
    insert_after = target
    for b in bullets:
        np = insert_paragraph_after(doc, insert_after, style)
        set_paragraph_text(np, f"{chr(8226)} {b}")
        insert_after = np

    delete_paragraph(target)


def normalize_competence_line(role: dict) -> str:
    """
    Return a canonical competence line in the format:
    skill1 | skill2 | skill3
    """
    raw = role.get("competence_line")
    if raw is None:
        raw = role.get("competencies", [])

    if isinstance(raw, (list, tuple)):
        parts = [str(x).strip() for x in raw if str(x).strip()]
        return " | ".join(parts)

    text = str(raw).strip()
    if not text:
        return ""

    if ":" in text and text.lower().startswith("kompetanse"):
        text = text.split(":", 1)[1].strip()

    parts = [p.strip() for p in re.split(r"\||,|;", text) if p.strip()]
    return " | ".join(parts)


# -----------------------------
# EXPERIENCE
# -----------------------------
def render_highlighted_experience(doc: Document, highlighted_roles: list):
    proto_h = proto_c = proto_b = proto_k = None

    for i, p in enumerate(doc.paragraphs):
        if "{{ROLE_TITLE}}" in p.text:
            proto_h = p
            proto_c = doc.paragraphs[i + 1]
            proto_b = doc.paragraphs[i + 2]
            proto_k = doc.paragraphs[i + 3]
            break

    if not proto_h:
        raise ValueError("Highlighted experience prototype not found in template.")

    header_style = proto_h.style
    body_style = proto_c.style
    insert_after = proto_k

    for idx, role in enumerate(highlighted_roles or []):
        ph = insert_paragraph_after(doc, insert_after, header_style)
        set_paragraph_text(
            ph,
            f"{role['title']} - {role['company']} ({role['period']})",
            bold=True,
        )

        pc = insert_paragraph_after(doc, ph, body_style)
        set_paragraph_text(
            pc,
            f"{role['client_context']}. {role['role_summary']}",
        )

        last = pc
        for b in role.get("bullets", []):
            pb = insert_paragraph_after(doc, last, body_style)
            set_paragraph_text(pb, f"{chr(8226)} {b}")
            last = pb

        pk = insert_paragraph_after(doc, last, body_style)
        set_paragraph_text(pk, normalize_competence_line(role))
        last = pk

        if idx != len(highlighted_roles) - 1:
            spacer = insert_paragraph_after(doc, last, body_style)
            set_paragraph_text(spacer, "")
            last = spacer

        insert_after = last

    delete_paragraph(proto_k)
    delete_paragraph(proto_b)
    delete_paragraph(proto_c)
    delete_paragraph(proto_h)


def render_other_experience(doc: Document, other_roles: list):
    proto_h = proto_c = proto_k = None

    for i, p in enumerate(doc.paragraphs):
        if "{{SHORT_ROLE_TITLE}}" in p.text:
            proto_h = p
            proto_c = doc.paragraphs[i + 1]
            proto_k = doc.paragraphs[i + 2]
            break

    if not proto_h:
        raise ValueError("Other experience prototype not found in template.")

    header_style = proto_h.style
    body_style = proto_c.style
    insert_after = proto_k

    for idx, role in enumerate(other_roles or []):
        ph = insert_paragraph_after(doc, insert_after, header_style)
        set_paragraph_text(
            ph,
            f"{role['title']} - {role['company']} ({role['period']})",
            bold=True,
        )

        pc = insert_paragraph_after(doc, ph, body_style)
        set_paragraph_text(
            pc,
            f"{role['client_context']}. {role['role_summary']}",
        )

        pk = insert_paragraph_after(doc, pc, body_style)
        set_paragraph_text(pk, normalize_competence_line(role))
        last = pk

        if idx != len(other_roles) - 1:
            spacer = insert_paragraph_after(doc, last, body_style)
            set_paragraph_text(spacer, "")
            last = spacer

        insert_after = last

    delete_paragraph(proto_k)
    delete_paragraph(proto_c)
    delete_paragraph(proto_h)


# -----------------------------
# PHOTO SWAP
# -----------------------------
def make_square_png_bytes(photo_path):
    img = Image.open(photo_path).convert("RGBA")
    w, h = img.size
    side = min(w, h)
    left = (w - side) // 2
    top = (h - side) // 2
    img_sq = img.crop((left, top, left + side, top + side)).resize((PHOTO_TARGET_PX, PHOTO_TARGET_PX))
    buf = io.BytesIO()
    img_sq.save(buf, format="PNG", optimize=True)
    return buf.getvalue()


def swap_profile_photo(docx_in, photo_path, docx_out):
    new_bytes = make_square_png_bytes(photo_path)

    with ZipFile(docx_in, "r") as zin, ZipFile(docx_out, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == PROFILE_PHOTO_MEDIA_PATH:
                data = new_bytes
            zout.writestr(item, data)


# -----------------------------
# PLACEHOLDER REPLACEMENT (ROBUST)
# -----------------------------
def _normalize_mapping_keys(mapping: dict) -> dict:
    """Ensure keys are always in {{KEY}} form, so we don't replace inside tokens."""
    normalized = {}
    for k, v in (mapping or {}).items():
        kk = str(k).strip()
        if not (kk.startswith("{{") and kk.endswith("}}")):
            kk = "{{" + kk + "}}"
        normalized[kk] = "" if v is None else str(v)
    return normalized


def _replace_in_paragraph(p: Paragraph, mapping: dict):
    for k, v in mapping.items():
        if k in p.text:
            if k == "{{NAME}}":
                set_paragraph_text(p, p.text.replace(k, v), bold=True, size_pt=16)
            else:
                set_paragraph_text(p, p.text.replace(k, v))


def replace_placeholders_everywhere(doc: Document, mapping: dict):
    mapping = _normalize_mapping_keys(mapping)

    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)

    for section in doc.sections:
        for p in section.header.paragraphs:
            _replace_in_paragraph(p, mapping)
        for p in section.footer.paragraphs:
            _replace_in_paragraph(p, mapping)


# -----------------------------
# MAIN GENERATOR
# -----------------------------
def generate_cv(
    template_path,
    candidate_photo_path,
    mapping,
    key_bullets,
    jd_bullets,
    highlighted_roles,
    other_roles,
    output_path,
):

    doc = Document(template_path)

    replace_placeholders_everywhere(doc, mapping)

    expand_bullets(doc, "{{KEY_ACHIEVEMENT_BULLETS}}", key_bullets or [])
    expand_bullets(doc, "{{JD_MATCH_BULLETS}}", jd_bullets or [])
    render_highlighted_experience(doc, highlighted_roles or [])
    render_other_experience(doc, other_roles or [])

    tmp = output_path.replace(".docx", "_tmp.docx")
    doc.save(tmp)
    swap_profile_photo(tmp, candidate_photo_path, output_path)
    os.remove(tmp)

    return output_path
